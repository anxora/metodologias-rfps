#!/usr/bin/env python3
"""
Generador simple de metodología para P2-P2RS Senegal
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def main():
    doc = Document()

    # Configurar márgenes
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # Título principal
    title = doc.add_heading("SECTION D : DESCRIPTION DE L'APPROCHE, MÉTHODOLOGIE ET PLAN DE TRAVAIL", level=1)

    # ============ INTRODUCTION ============
    doc.add_heading("Introduction", level=2)
    doc.add_paragraph("""Notre consortium propose une approche méthodologique intégrée et participative pour l'accompagnement des communautés des 15 communes cibles du P2-P2RS au Sénégal. Forts de notre expérience dans les projets de gestion des ressources naturelles au Sahel, nous avons conçu une méthodologie qui s'ancre dans les réalités locales tout en répondant aux exigences de la Banque Africaine de Développement.

Notre approche repose sur trois piliers fondamentaux : la participation communautaire, le renforcement des capacités locales et la durabilité des interventions. Nous mobiliserons une équipe pluridisciplinaire expérimentée dans les régions de Matam, Tambacounda et Fatick pour assurer un accompagnement de proximité efficace des Comités de Gestion Concertée des Ressources Naturelles (CGC/RN) et la mise en œuvre réussie du Mécanisme de Gestion des Plaintes (MGP).""")

    # ============ COMPRÉHENSION DU CONTEXTE ============
    doc.add_heading("Compréhension du Contexte", level=2)

    doc.add_heading("Contexte du Programme P2RS", level=3)
    doc.add_paragraph("""Le Programme régional de résilience à l'insécurité alimentaire et nutritionnelle au Sahel (P2RS), initié par la Banque Africaine de Développement depuis 2014, constitue une réponse structurelle aux défis chroniques de sécurité alimentaire dans la région sahélienne. Le P2-P2RS (Projet 2) s'inscrit dans la continuité des acquis de la phase 1, avec un accent particulier sur la consolidation et l'élargissement des interventions dans 15 communes réparties dans les régions de Matam, Tambacounda et Fatick.""")

    doc.add_heading("Zones d'intervention et enjeux spécifiques", level=3)
    doc.add_paragraph("""Le projet intervient dans deux pôles de développement distincts aux caractéristiques contrastées :

Le Pôle Nord (Vallée du fleuve Sénégal) regroupe 12 communes dans les départements de Matam, Kanel et Bakel. Cette zone agro-écologique se caractérise par des systèmes agropastoraux transhumants, une forte dépendance aux ressources hydriques du fleuve Sénégal, et des défis liés à la dégradation des terres et aux conflits agriculteurs-éleveurs.

Le Pôle Centre (Bassin arachidier) comprend 3 communes insulaires (Fimela, Bassoul, Djirnda) dans la région de Fatick, particulièrement vulnérables aux effets des changements climatiques, notamment la salinisation des terres et l'élévation du niveau de la mer.""")

    doc.add_heading("Cadre institutionnel et parties prenantes", level=3)
    doc.add_paragraph("""Le projet s'inscrit dans le cadre de la politique nationale de décentralisation du Sénégal, avec les collectivités territoriales comme porte d'entrée principale. Les acteurs clés incluent : les services techniques déconcentrés (DRDR, IREF, Service de l'élevage), les organisations de producteurs agricoles et pastoraux, les groupements de femmes, les autorités coutumières et religieuses, ainsi que les ONG locales actives dans la gestion des ressources naturelles.""")

    doc.add_heading("Problématique et défis majeurs", level=3)
    doc.add_paragraph("""Les communautés cibles font face à plusieurs défis interconnectés :""")

    for item in [
        "La dégradation continue des terres agricoles due à l'érosion hydrique et éolienne, la salinisation et la perte de fertilité",
        "Les conflits récurrents entre agriculteurs et éleveurs liés à la réduction des espaces pastoraux et à l'obstruction des couloirs de passage",
        "La faible capacité organisationnelle des structures communautaires de gestion des ressources naturelles",
        "L'insuffisance des connaissances en techniques de Gestion Durable des Terres (GDT) et d'agroécologie",
        "L'absence de mécanismes efficaces de résolution des conflits et de gestion des plaintes"
    ]:
        doc.add_paragraph(item, style='List Bullet')

    # ============ PRINCIPES DIRECTEURS ============
    doc.add_heading("Nos Principes Directeurs", level=2)
    doc.add_paragraph("""Les principes clés suivants sous-tendent notre stratégie proposée et découlent de notre analyse du contexte local ainsi que des besoins spécifiques exprimés dans les TdR.""")

    principles = [
        ("Principe 1 : Approche participative et inclusive",
         "Toutes nos interventions seront fondées sur la participation effective des communautés à toutes les étapes : diagnostic, planification, mise en œuvre et suivi. Nous accorderons une attention particulière à l'inclusion des groupes vulnérables, notamment les femmes et les jeunes, dans les instances de décision et les activités de renforcement des capacités."),

        ("Principe 2 : Ancrage dans les savoirs locaux",
         "Notre méthodologie valorise et capitalise sur les connaissances endogènes des communautés en matière de gestion des ressources naturelles. Les pratiques traditionnelles validées seront intégrées aux innovations techniques proposées."),

        ("Principe 3 : Adaptation au contexte local",
         "Reconnaissant la diversité des contextes agro-écologiques entre le Pôle Nord et le Pôle Centre, nous adapterons nos méthodes et contenus de formation aux spécificités de chaque zone. Les outils de communication seront développés dans les langues locales (Pulaar, Wolof, Sérère)."),

        ("Principe 4 : Apprentissage par l'action",
         "Notre programme de renforcement des capacités privilégie l'apprentissage pratique à travers les Champs Écoles Paysans et les démonstrations sur le terrain."),

        ("Principe 5 : Coordination et synergie",
         "Nous assurerons une coordination étroite avec les services techniques déconcentrés, les autres partenaires du P2RS et les projets connexes intervenant dans les mêmes zones."),

        ("Principe 6 : Équité genre et autonomisation",
         "L'intégration du genre sera transversale à toutes nos interventions. Nous veillerons à ce que les femmes représentent au moins 40% des participants aux formations et au moins 30% des membres des instances de décision des CGC/RN.")
    ]

    for name, desc in principles:
        p = doc.add_paragraph()
        run = p.add_run(name)
        run.bold = True
        doc.add_paragraph(desc)

    # ============ APPROCHE TECHNIQUE ET MÉTHODOLOGIE ============
    doc.add_heading("Approche Technique et Méthodologie", level=1)

    # PHASE 1
    doc.add_heading("Phase 1 : Démarrage et Diagnostic Participatif (Mois 1-4)", level=2)
    doc.add_paragraph("""Cette phase fondamentale permettra d'établir les bases solides de notre intervention à travers une compréhension approfondie du contexte socioéconomique et des dynamiques de gestion des ressources naturelles dans les 15 communes cibles. Nous procéderons à l'installation de nos équipes dans les trois régions et établirons les premiers contacts avec les autorités locales et les communautés bénéficiaires.""")

    doc.add_heading("Activité 1A : Installation et cadrage méthodologique", level=3)
    doc.add_paragraph("""L'activité de démarrage comprendra l'installation des bases opérationnelles dans les trois régions d'intervention. Nous procéderons au recrutement et à la formation de l'équipe locale (animateurs), à la mise en place des moyens logistiques et à l'établissement des protocoles de travail avec l'Unité de Coordination du Projet.""")
    for item in [
        "Installation des bases opérationnelles à Matam, Bakel et Fatick",
        "Recrutement et formation des 13 animateurs locaux",
        "Acquisition et déploiement des moyens logistiques",
        "Atelier de lancement avec les parties prenantes du P2RS",
        "Visites de courtoisie aux autorités des 15 communes"
    ]:
        doc.add_paragraph(item, style='List Bullet')
    p = doc.add_paragraph()
    run = p.add_run("Livrable L1 : Rapport d'établissement avec recadrage méthodologique et planning actualisé")
    run.bold = True

    doc.add_heading("Activité 1B : Études diagnostiques socioéconomiques", level=3)
    doc.add_paragraph("""Le diagnostic socioéconomique sera conduit dans chacun des villages ciblés selon une méthodologie participative combinant : enquêtes ménages, focus groups différenciés (hommes, femmes, jeunes, éleveurs), entretiens avec les personnes ressources et observations directes.""")
    for item in [
        "Élaboration des outils de collecte de données (questionnaires, guides d'entretien)",
        "Formation des enquêteurs sur les méthodologies participatives",
        "Réalisation des enquêtes ménages et focus groups dans les 15 communes",
        "Cartographie participative des ressources naturelles",
        "Analyse des données et rédaction des monographies communales"
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("Activité 1C : Inventaire des stratégies endogènes de GRN", level=3)
    doc.add_paragraph("""Cette activité vise à documenter et valoriser les pratiques traditionnelles de gestion des ressources naturelles qui ont fait leurs preuves dans les communautés.""")
    for item in [
        "Recension des pratiques traditionnelles de GRN par zone agro-écologique",
        "Documentation des conventions locales et règles coutumières existantes",
        "Identification des personnes ressources et détenteurs de savoirs locaux",
        "Analyse comparative des stratégies endogènes efficaces",
        "Élaboration d'un répertoire des bonnes pratiques locales"
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("Activité 1D : Restitution et validation du diagnostic", level=3)
    doc.add_paragraph("""Les résultats du diagnostic seront restitués et validés lors d'ateliers participatifs organisés au niveau de chaque commune et au niveau régional.""")
    for item in [
        "Organisation de 15 ateliers communaux de restitution",
        "Organisation de 3 ateliers régionaux de validation",
        "Intégration des observations et corrections",
        "Finalisation des rapports de diagnostic par région",
        "Diffusion des résultats aux parties prenantes"
    ]:
        doc.add_paragraph(item, style='List Bullet')
    p = doc.add_paragraph()
    run = p.add_run("Livrable L2 : Rapports de diagnostic socioéconomique et environnemental des 15 communes")
    run.bold = True

    # PHASE 2
    doc.add_heading("Phase 2 : Mise en Place et Opérationnalisation des CGC/RN (Mois 5-10)", level=2)
    doc.add_paragraph("""Cette phase constitue le cœur de notre intervention pour l'établissement d'un système durable de gestion concertée des ressources naturelles. Sur la base des résultats du diagnostic, nous procéderons à la mise en place ou à la redynamisation des Comités de Gestion Concertée des Ressources Naturelles (CGC/RN) dans chaque commune.""")

    doc.add_heading("Activité 2A : Sensibilisation et mobilisation communautaire", level=3)
    doc.add_paragraph("""Une vaste campagne d'information et de sensibilisation sera déployée dans les 15 communes pour expliquer les objectifs du projet et l'importance de la gestion concertée des ressources naturelles.""")
    for item in [
        "Élaboration de la stratégie et des supports de communication",
        "Production d'émissions radiophoniques en Pulaar, Wolof et Sérère",
        "Organisation de réunions d'information dans chaque village",
        "Formation de relais communautaires pour la sensibilisation continue",
        "Mise en place de panneaux d'information dans les communes"
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("Activité 2B : Constitution et structuration des CGC/RN", level=3)
    doc.add_paragraph("""Les CGC/RN seront mis en place ou redynamisés selon un processus transparent et participatif. Chaque CGC/RN comprendra des représentants des agriculteurs, éleveurs, exploitants forestiers, femmes, jeunes et autorités coutumières.""")
    for item in [
        "Évaluation des structures existantes de gestion des RN",
        "Facilitation des assemblées générales constitutives",
        "Élection des bureaux des CGC/RN (avec quota genre)",
        "Élaboration des statuts et règlements intérieurs",
        "Reconnaissance officielle des CGC/RN par les autorités"
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("Activité 2C : Négociation et élaboration des conventions locales", level=3)
    doc.add_paragraph("""Les règles de gestion concertée des ressources naturelles seront négociées à travers des ateliers multi-acteurs regroupant tous les groupes d'utilisateurs.""")
    for item in [
        "Organisation d'ateliers de négociation multi-acteurs par commune",
        "Facilitation des consensus sur les règles de gestion",
        "Rédaction des projets de conventions locales",
        "Validation participative des conventions locales",
        "Adoption des conventions par les conseils municipaux"
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("Activité 2D : Délimitation et matérialisation des espaces", level=3)
    doc.add_paragraph("""En lien avec les POAS existants, nous appuierons la délimitation et la matérialisation physique des espaces convenus : couloirs de passage du bétail, zones de pâturage, aires de repos, zones de mise en défens.""")
    for item in [
        "Cartographie participative des espaces par commune",
        "Identification et bornage des couloirs de passage",
        "Matérialisation des zones de pâturage et aires de repos",
        "Délimitation des zones de mise en défens",
        "Production et diffusion des cartes des terroirs"
    ]:
        doc.add_paragraph(item, style='List Bullet')
    p = doc.add_paragraph()
    run = p.add_run("Livrable L3 : 15 CGC/RN opérationnels avec conventions locales adoptées et espaces délimités")
    run.bold = True

    # PHASE 3
    doc.add_heading("Phase 3 : Formation et Renforcement des Capacités (Mois 6-14)", level=2)
    doc.add_paragraph("""Cette phase est dédiée au renforcement des capacités des acteurs à tous les niveaux. Le programme de formation couvrira les thématiques de la Gestion Durable des Terres (GDT), de l'agroécologie, de l'agroforesterie, de la Régénération Naturelle Assistée (RNA) et de la gestion des conflits.""")

    doc.add_heading("Activité 3A : Élaboration du plan de renforcement des capacités", level=3)
    doc.add_paragraph("""Sur la base des besoins identifiés lors du diagnostic, nous élaborerons un plan détaillé de renforcement des capacités définissant les thématiques prioritaires, les approches pédagogiques, les contenus des modules et le calendrier des sessions.""")
    for item in [
        "Analyse des besoins de formation par groupe cible",
        "Définition des objectifs et contenus pédagogiques",
        "Identification des formateurs et personnes ressources",
        "Élaboration du calendrier de formation",
        "Validation du plan par l'UCP et les partenaires"
    ]:
        doc.add_paragraph(item, style='List Bullet')
    p = doc.add_paragraph()
    run = p.add_run("Livrable L4 : Plan de renforcement des capacités validé avec modules de formation")
    run.bold = True

    doc.add_heading("Activité 3B : Formation des formateurs et animateurs", level=3)
    doc.add_paragraph("""Des sessions de formation des formateurs seront organisées pour nos techniciens et animateurs ainsi que les agents des services techniques partenaires.""")
    for item in [
        "Session de formation sur les techniques de facilitation (5 jours)",
        "Session technique sur la GDT et CES/DRS (5 jours)",
        "Session technique sur l'agroécologie et la RNA (5 jours)",
        "Session sur la production de plants agroforestiers (3 jours)",
        "Session sur la gestion des conflits et médiation (3 jours)"
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("Activité 3C : Formation des membres des CGC/RN", level=3)
    doc.add_paragraph("""Les membres des 15 CGC/RN seront formés sur leurs rôles et responsabilités, les techniques de gouvernance locale, la gestion des conflits et le suivi-évaluation participatif.""")
    for item in [
        "Module gouvernance et leadership (2 jours x 15 CGC/RN)",
        "Module techniques de GDT (3 jours x 15 CGC/RN)",
        "Module gestion des conflits (2 jours x 15 CGC/RN)",
        "Module suivi-évaluation participatif (2 jours x 15 CGC/RN)",
        "Visite d'échange inter-CGC/RN (1 jour)"
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("Activité 3D : Formation des producteurs via Champs Écoles", level=3)
    doc.add_paragraph("""L'approche Champs Écoles Paysans (CEP) sera déployée pour former les producteurs aux techniques de GDT, d'agroécologie et d'agroforesterie. Chaque CEP regroupera 25 à 30 producteurs.""")
    for item in [
        "Installation de 30 Champs Écoles Paysans (2 par commune)",
        "Sessions pratiques sur le compostage et la fertilité intégrée",
        "Sessions pratiques sur les ouvrages antiérosifs (cordons pierreux, zaï)",
        "Sessions pratiques sur la RNA et l'agroforesterie",
        "Évaluation participative des apprentissages"
    ]:
        doc.add_paragraph(item, style='List Bullet')
    p = doc.add_paragraph()
    run = p.add_run("Livrable L5 : Rapport de formation avec 450 membres CGC/RN et 900 producteurs formés")
    run.bold = True

    # PHASE 4
    doc.add_heading("Phase 4 : Accompagnement et Mise en Œuvre du MGP (Mois 5-24)", level=2)
    doc.add_paragraph("""Cette phase transversale est dédiée à l'accompagnement continu des CGC/RN dans la mise en œuvre des conventions locales et l'opérationnalisation du Mécanisme de Gestion des Plaintes (MGP).""")

    doc.add_heading("Activité 4A : Installation du dispositif MGP", level=3)
    doc.add_paragraph("""Le Mécanisme de Gestion des Plaintes sera installé à trois niveaux : comités villageois, comités communaux et comité régional.""")
    for item in [
        "Identification des membres des comités de gestion des plaintes",
        "Organisation des assemblées d'installation des comités",
        "Dotation en matériels et outils de travail",
        "Mise en place des points de dépôt et du numéro vert",
        "Élaboration des procédures et circuits de traitement"
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("Activité 4B : Formation des membres des comités MGP", level=3)
    doc.add_paragraph("""Les membres des comités de gestion des plaintes seront formés sur la politique de gestion des plaintes, les procédures de réception et traitement, et les techniques de médiation.""")
    for item in [
        "Formation sur la politique et les procédures du MGP (2 jours)",
        "Formation sur les techniques de médiation (2 jours)",
        "Formation sur la documentation et le rapportage (1 jour)",
        "Sessions de simulation et mise en situation",
        "Recyclage périodique des membres"
    ]:
        doc.add_paragraph(item, style='List Bullet')
    p = doc.add_paragraph()
    run = p.add_run("Livrable L6 : Dispositif MGP opérationnel avec comités formés et outils déployés")
    run.bold = True

    doc.add_heading("Activité 4C : Vulgarisation du MGP auprès des communautés", level=3)
    doc.add_paragraph("""Une campagne intensive de vulgarisation sera déployée pour informer les populations sur l'existence du MGP et les canaux de dépôt disponibles.""")
    for item in [
        "Production de spots radio sur le MGP en langues locales",
        "Conception et distribution d'affiches et dépliants illustrés",
        "Organisation de réunions d'information dans tous les villages",
        "Formation de relais communautaires pour la sensibilisation",
        "Rappels périodiques lors des activités du projet"
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("Activité 4D : Accompagnement des CGC/RN et suivi des conventions", level=3)
    doc.add_paragraph("""Nos animateurs assureront un accompagnement régulier des CGC/RN dans l'exercice de leurs fonctions : animation des réunions, application des conventions locales, gestion des infractions.""")
    for item in [
        "Visites hebdomadaires d'accompagnement des CGC/RN",
        "Appui à l'animation des réunions statutaires",
        "Suivi de l'application des conventions locales",
        "Facilitation de la résolution des conflits",
        "Appui à la mobilisation des ressources locales"
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("Activité 4E : Suivi-évaluation et capitalisation", level=3)
    doc.add_paragraph("""Un dispositif de suivi-évaluation participatif sera mis en place pour mesurer les progrès et documenter les leçons apprises.""")
    for item in [
        "Mise en place du système de suivi-évaluation",
        "Collecte et analyse des données de suivi",
        "Élaboration des rapports trimestriels d'activités",
        "Évaluations participatives avec les communautés",
        "Capitalisation et documentation des bonnes pratiques"
    ]:
        doc.add_paragraph(item, style='List Bullet')
    p = doc.add_paragraph()
    run = p.add_run("Livrable L7 : 8 rapports trimestriels d'activités et rapport final de capitalisation")
    run.bold = True

    # ============ GESTION DES RISQUES ============
    doc.add_heading("Gestion des Risques", level=1)

    doc.add_heading("Risques contextuels", level=3)
    doc.add_paragraph("""• Instabilité sécuritaire : La région nord peut connaître des tensions liées aux mouvements transfrontaliers. Mesure d'atténuation : coordination étroite avec les autorités locales.

• Conditions climatiques défavorables : Les aléas climatiques peuvent perturber les activités agricoles. Mesure d'atténuation : flexibilité du calendrier et priorisation des activités critiques.

• Conflits intercommunautaires : Les tensions agriculteurs-éleveurs peuvent s'intensifier. Mesure d'atténuation : renforcement de la médiation préventive.""")

    doc.add_heading("Risques opérationnels", level=3)
    doc.add_paragraph("""• Mobilisation insuffisante des communautés : Mesure d'atténuation : approche participative dès le diagnostic, communication intensive.

• Turnover du personnel : Mesure d'atténuation : documentation systématique, partage des connaissances.

• Difficultés logistiques : Mesure d'atténuation : planification anticipée, moyens adaptés (pirogues pour les îles).""")

    doc.add_heading("Risques institutionnels", level=3)
    doc.add_paragraph("""• Faible appropriation par les autorités locales : Mesure d'atténuation : implication dès le départ des maires et conseils municipaux.

• Conflits de compétences : Mesure d'atténuation : cartographie des acteurs, coordination régulière.""")

    # ============ ASSURANCE QUALITÉ ============
    doc.add_heading("Assurance Qualité", level=1)

    doc.add_heading("Indicateurs clés de performance (KPIs)", level=3)
    for item in [
        "Taux de fonctionnement des CGC/RN : cible 80%",
        "Taux de résolution des plaintes dans les délais : cible 90%",
        "Taux de satisfaction des bénéficiaires des formations : cible 85%",
        "Représentation des femmes dans les instances de décision : minimum 30%",
        "Taux d'application des conventions locales : cible 75%"
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("Mécanismes de contrôle et supervision", level=3)
    for item in [
        "Supervision mensuelle du chef de mission dans chaque région",
        "Réunions hebdomadaires de coordination des techniciens",
        "Visites de terrain conjointes avec l'UCP (trimestrielles)",
        "Audits externes annuels des activités",
        "Évaluations participatives semestrielles avec les communautés"
    ]:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("Protocoles de reporting", level=3)
    for item in [
        "Rapports hebdomadaires des animateurs (format standardisé)",
        "Rapports mensuels des techniciens avec données de suivi",
        "Rapports trimestriels consolidés à l'UCP (5 jours après fin trimestre)",
        "Base de données géoréférencée des interventions"
    ]:
        doc.add_paragraph(item, style='List Bullet')

    # ============ PLAN DE TRAVAIL ============
    doc.add_heading("Plan de Travail", level=1)

    # Crear tabla simple del plan de trabajo
    table = doc.add_table(rows=9, cols=5)
    table.style = 'Table Grid'

    # Headers
    headers = ['Phase/Activité', 'M1-4', 'M5-10', 'M11-18', 'M19-24']
    for i, header in enumerate(headers):
        table.cell(0, i).text = header

    # Datos
    data = [
        ['Phase 1: Démarrage et Diagnostic', 'X', '', '', ''],
        ['Phase 2: Mise en place CGC/RN', '', 'X', '', ''],
        ['Phase 3: Formation et Renforcement', '', 'X', 'X', ''],
        ['Phase 4: Accompagnement et MGP', '', 'X', 'X', 'X'],
        ['', '', '', '', ''],
        ['Livrables:', '', '', '', ''],
        ['L1-L2: Rapports diagnostic', 'X', '', '', ''],
        ['L3-L6: CGC/RN, Formations, MGP', '', 'X', 'X', ''],
    ]

    for i, row_data in enumerate(data):
        for j, cell_data in enumerate(row_data):
            table.cell(i+1, j).text = cell_data

    # Guardar
    output_path = "/Users/apple/metodologias-rfps/output/Methodologie_P2RS_Senegal.docx"
    doc.save(output_path)
    print(f"Documento generado: {output_path}")


if __name__ == "__main__":
    main()
