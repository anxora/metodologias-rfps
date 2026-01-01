"""
Generador de documentos Word con formato Aninver
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# Colores Aninver
COLORS = {
    'principle_header': 'BDD6EE',      # Azul claro - header de principios
    'principle_title': '2E74B5',        # Azul oscuro - título de principio
    'activities_header': '5B83DF',      # Azul - header "Activities"
    'activities_body': 'F2F2F2',        # Gris claro - lista de actividades
    'timeline': '00A798',               # Verde/Teal - "Timeline"
    'gantt_active': 'FFC000',           # Amarillo/Naranja - celdas activas del Gantt
    'experience_box': 'FFF2CC',         # Amarillo claro - cajas de experiencia
    'deliverable_header': 'FFC000',     # Amarillo - header de entregable
    'deliverable_body': 'FFF2CC',       # Amarillo claro - cuerpo de entregable
}


def create_word_document(methodology: dict, output_path: str, lang_config: dict):
    """
    Crea un documento Word con el enfoque metodológico en formato Aninver
    """
    doc = Document()
    sections = lang_config['sections']

    # Configurar estilos Aninver
    _setup_aninver_styles(doc)

    # Configurar márgenes
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # Título principal de sección
    title = doc.add_paragraph(sections['title'], style='Aninver Section Heading')

    # Introducción si existe
    if methodology.get('introduction'):
        _add_paragraph(doc, methodology['introduction'])

    # Comprensión del contexto
    if methodology.get('context'):
        doc.add_paragraph(sections['context'], style='Aninver Title 2')
        _add_content_block(doc, methodology['context'])

    # Principios rectores
    if methodology.get('principles'):
        doc.add_paragraph(sections.get('principles', 'Our guiding principles'), style='Aninver Title 2')
        intro_text = lang_config.get('principles_intro',
            'The following key principles underpin our proposed strategy, and stem from our analysis of the local context as well as the specific needs expressed in the ToR.')
        _add_paragraph(doc, intro_text)
        _add_principles_table(doc, methodology['principles'])

    # Enfoque técnico y metodología
    doc.add_paragraph(sections['approach'], style='Aninver Title 1')

    # Fases del proyecto
    if methodology.get('phases'):
        for phase in methodology['phases']:
            _add_phase_improved(doc, phase, lang_config)

    # Gestión de riesgos
    if methodology.get('risks'):
        doc.add_paragraph(sections['risks'], style='Aninver Title 1')
        _add_content_block(doc, methodology['risks'])

    # Aseguramiento de calidad
    if methodology.get('quality'):
        doc.add_paragraph(sections['quality'], style='Aninver Title 1')
        _add_content_block(doc, methodology['quality'])

    # Guardar documento
    doc.save(output_path)


def _setup_aninver_styles(doc: Document):
    """Configura los estilos Aninver del documento"""
    styles = doc.styles

    # Configuración base del documento: Arial 10, interlineado Múltiple 1.3, anterior 0, posterior 6
    FONT_NAME = 'Arial'
    FONT_SIZE = Pt(10)
    LINE_SPACING = 1.3  # Múltiple 1.3
    SPACE_BEFORE = Pt(0)
    SPACE_AFTER = Pt(6)

    # Estilo base Normal
    style_normal = styles['Normal']
    style_normal.font.name = FONT_NAME
    style_normal.font.size = FONT_SIZE
    style_normal.paragraph_format.line_spacing = LINE_SPACING
    style_normal.paragraph_format.space_before = SPACE_BEFORE
    style_normal.paragraph_format.space_after = SPACE_AFTER
    style_normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Aninver Section Heading (título principal de sección)
    try:
        style = styles.add_style('Aninver Section Heading', WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        style = styles['Aninver Section Heading']
    style.font.name = FONT_NAME
    style.font.size = Pt(14)
    style.font.bold = True
    style.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    style.paragraph_format.space_before = Pt(12)
    style.paragraph_format.space_after = SPACE_AFTER
    style.paragraph_format.line_spacing = LINE_SPACING

    # Aninver Title 1 (títulos principales)
    try:
        style = styles.add_style('Aninver Title 1', WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        style = styles['Aninver Title 1']
    style.font.name = FONT_NAME
    style.font.size = Pt(12)
    style.font.bold = True
    style.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    style.paragraph_format.space_before = Pt(12)
    style.paragraph_format.space_after = SPACE_AFTER
    style.paragraph_format.line_spacing = LINE_SPACING

    # Aninver Title 2 (subtítulos de fase)
    try:
        style = styles.add_style('Aninver Title 2', WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        style = styles['Aninver Title 2']
    style.font.name = FONT_NAME
    style.font.size = Pt(11)
    style.font.bold = True
    style.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    style.paragraph_format.space_before = Pt(6)
    style.paragraph_format.space_after = SPACE_AFTER
    style.paragraph_format.line_spacing = LINE_SPACING

    # Aninver Title 3 (subtítulos de actividad)
    try:
        style = styles.add_style('Aninver Title 3', WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        style = styles['Aninver Title 3']
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE
    style.font.bold = True
    style.font.italic = True
    style.paragraph_format.space_before = SPACE_BEFORE
    style.paragraph_format.space_after = SPACE_AFTER
    style.paragraph_format.line_spacing = LINE_SPACING

    # Aninver Body Text
    try:
        style = styles.add_style('Aninver Body Text', WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        style = styles['Aninver Body Text']
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE
    style.paragraph_format.space_before = SPACE_BEFORE
    style.paragraph_format.space_after = SPACE_AFTER
    style.paragraph_format.line_spacing = LINE_SPACING
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Aninver Bullet List
    try:
        style = styles.add_style('Aninver Bullet List', WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        style = styles['Aninver Bullet List']
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE
    style.paragraph_format.left_indent = Inches(0.5)
    style.paragraph_format.space_before = SPACE_BEFORE
    style.paragraph_format.space_after = SPACE_AFTER
    style.paragraph_format.line_spacing = LINE_SPACING
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def _add_paragraph(doc: Document, text: str):
    """Agrega un párrafo con estilo Aninver Body Text, procesando negritas **texto**"""
    if not text:
        return None

    p = doc.add_paragraph(style='Aninver Body Text')

    # Procesar negritas marcadas con **texto**
    import re
    parts = re.split(r'(\*\*[^*]+\*\*)', text)

    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Texto en negrita
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            # Texto normal
            if part:
                p.add_run(part)

    return p


def _add_content_block(doc: Document, content: str):
    """Agrega un bloque de contenido, detectando listas y párrafos, procesando negritas"""
    import re

    if not content:
        return

    paragraphs = content.split('\n\n')
    for para_text in paragraphs:
        if not para_text.strip():
            continue

        lines = para_text.strip().split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Detectar headers con **texto** (línea completa en negrita)
            if line.startswith('**') and line.endswith('**') and line.count('**') == 2:
                p = doc.add_paragraph()
                run = p.add_run(line[2:-2])
                run.bold = True
                p.paragraph_format.space_before = Pt(8)
                continue

            if line.startswith('- ') or line.startswith('• ') or line.startswith('* '):
                p = doc.add_paragraph(style='List Bullet')
                text = line[2:]
                _add_text_with_bold(p, text)
            elif line.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
                text = line.split('.', 1)[1].strip() if '.' in line else line
                p = doc.add_paragraph(style='List Number')
                _add_text_with_bold(p, text)
            else:
                _add_paragraph(doc, line)


def _add_text_with_bold(paragraph, text: str):
    """Agrega texto a un párrafo procesando negritas **texto**"""
    import re
    parts = re.split(r'(\*\*[^*]+\*\*)', text)

    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            if part:
                paragraph.add_run(part)


def _set_cell_shading(cell, color_hex: str):
    """Establece el color de fondo de una celda"""
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def _format_table_cell(cell, font_name: str = 'Aptos', font_size: int = 10):
    """Aplica formato Aptos 10 e interlineado sencillo a una celda de tabla"""
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)


def _format_all_table_cells(table, font_name: str = 'Aptos', font_size: int = 10):
    """Aplica formato a todas las celdas de una tabla"""
    for row in table.rows:
        for cell in row.cells:
            _format_table_cell(cell, font_name, font_size)


def _set_table_borders(table, color_hex: str = "000000", size: int = 4):
    """
    Establece bordes negros en todas las celdas de la tabla
    size: tamaño del borde en eighths of a point (4 = 0.5pt, 8 = 1pt)
    """
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')

    # Crear bordes de tabla
    tblBorders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="{size}" w:color="{color_hex}"/>'
        f'<w:left w:val="single" w:sz="{size}" w:color="{color_hex}"/>'
        f'<w:bottom w:val="single" w:sz="{size}" w:color="{color_hex}"/>'
        f'<w:right w:val="single" w:sz="{size}" w:color="{color_hex}"/>'
        f'<w:insideH w:val="single" w:sz="{size}" w:color="{color_hex}"/>'
        f'<w:insideV w:val="single" w:sz="{size}" w:color="{color_hex}"/>'
        f'</w:tblBorders>'
    )

    # Eliminar bordes existentes si los hay
    existing_borders = tblPr.find(qn('w:tblBorders'))
    if existing_borders is not None:
        tblPr.remove(existing_borders)

    tblPr.append(tblBorders)

    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def _remove_table_borders(table):
    """Elimina todos los bordes de una tabla"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')

    # Crear bordes invisibles
    tblBorders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )

    # Eliminar bordes existentes si los hay
    existing_borders = tblPr.find(qn('w:tblBorders'))
    if existing_borders is not None:
        tblPr.remove(existing_borders)

    tblPr.append(tblBorders)

    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def _add_principles_table(doc: Document, principles: list):
    """
    Agrega una tabla de principios con 5 columnas:
    - Columnas 0, 2, 4: Contenido (5.2 cm cada una)
    - Columnas 1, 3: Separadores de 1 cm
    - Sin bordes
    - Fila 0: Headers "Principle 1/2/3" en azul claro (BDD6EE)
    - Fila 1: Nombres de principios 1-3 en azul oscuro (2E74B5), texto blanco
    - Fila 2: Descripciones 1-3
    - Fila 3: Nombres de principios 4-6 en azul oscuro (sin header)
    - Fila 4: Descripciones 4-6
    """
    if not principles:
        return

    num_principles = len(principles)

    # Determinar número de filas: 3 para <=3 principios, 5 para >3 principios
    num_rows = 3 if num_principles <= 3 else 5

    # Crear tabla con 5 columnas (3 contenido + 2 separadores)
    table = doc.add_table(rows=num_rows, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False  # Desactivar autofit para controlar anchos

    # IMPORTANTE: Eliminar bordes de la tabla
    _remove_table_borders(table)

    # Anchos de columna
    separator_width = Cm(1)   # 1 cm para separadores
    content_width = Cm(5.2)   # 5.2 cm cada columna de contenido

    # Definir anchos por índice de columna
    col_widths = [content_width, separator_width, content_width, separator_width, content_width]

    # Aplicar anchos a cada celda individualmente (python-docx requiere esto)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = col_widths[idx]

    # Mapeo de columnas de contenido: 0, 2, 4 (saltando separadores 1 y 3)
    content_cols = [0, 2, 4]

    # === FILA 0: Headers "Principle 1", "Principle 2", "Principle 3" ===
    row0 = table.rows[0]
    for i, col_idx in enumerate(content_cols[:min(num_principles, 3)]):
        cell = row0.cells[col_idx]
        cell.text = f"Principle {i + 1}"
        _set_cell_shading(cell, COLORS['principle_header'])
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.name = 'Aptos'
                run.font.size = Pt(10)

    # === FILA 1: Nombres de principios 1-3 ===
    row1 = table.rows[1]
    for i, principle in enumerate(principles[:3]):
        col_idx = content_cols[i]
        cell = row1.cells[col_idx]
        cell.text = principle.get('name', '')
        _set_cell_shading(cell, COLORS['principle_title'])
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.name = 'Aptos'
                run.font.size = Pt(10)

    # === FILA 2: Descripciones de principios 1-3 ===
    row2 = table.rows[2]
    for i, principle in enumerate(principles[:3]):
        col_idx = content_cols[i]
        cell = row2.cells[col_idx]
        cell.text = principle.get('description', '')
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in para.runs:
                run.font.name = 'Aptos'
                run.font.size = Pt(10)

    # === Si hay más de 3 principios ===
    if num_principles > 3:
        # FILA 3: Nombres de principios 4-6 (sin header "Principle X")
        row3 = table.rows[3]
        for i, principle in enumerate(principles[3:6]):
            col_idx = content_cols[i]
            cell = row3.cells[col_idx]
            cell.text = principle.get('name', '')
            _set_cell_shading(cell, COLORS['principle_title'])
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.bold = True
                    run.font.name = 'Aptos'
                    run.font.size = Pt(10)

        # FILA 4: Descripciones de principios 4-6
        row4 = table.rows[4]
        for i, principle in enumerate(principles[3:6]):
            col_idx = content_cols[i]
            cell = row4.cells[col_idx]
            cell.text = principle.get('description', '')
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in para.runs:
                    run.font.name = 'Aptos'
                    run.font.size = Pt(10)

    # Aplicar interlineado sencillo a todas las celdas
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.line_spacing = 1.0
                para.paragraph_format.space_before = Pt(3)
                para.paragraph_format.space_after = Pt(3)

    doc.add_paragraph()


def _add_phase_improved(doc: Document, phase: dict, lang_config: dict):
    """
    Agrega una fase con el formato mejorado de tablas de actividades
    """
    # Título de la fase
    doc.add_paragraph(phase.get('title', ''), style='Aninver Title 2')

    # Descripción de la fase
    if phase.get('description'):
        _add_paragraph(doc, phase['description'])

    tasks = phase.get('tasks', [])
    if tasks:
        # 1. Tabla de Activities/Timeline (vertical)
        _add_activities_timeline_table_improved(doc, tasks, phase)

        # 2. Tabla Gantt con semanas/meses
        _add_gantt_table_improved(doc, tasks, phase)

        # 3. Detalle de cada tarea
        for task in tasks:
            task_code = task.get('code', '')
            task_title = task.get('title', '')
            doc.add_paragraph(f"{task_code}. {task_title}", style='Aninver Title 3')

            # Descripción de la tarea (sin forzar bullets)
            if task.get('description'):
                _add_content_block(doc, task['description'])

            # Items solo si existen y son relevantes
            if task.get('items'):
                for item in task['items']:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(item)

    # Entregables de la fase
    if phase.get('deliverables'):
        for deliverable in phase['deliverables']:
            _add_deliverable_box(doc, deliverable)

    doc.add_paragraph()


def _add_activities_timeline_table_improved(doc: Document, tasks: list, phase: dict):
    """
    Agrega la tabla vertical de Activities/Timeline con formato mejorado
    - Header "Activities" en azul
    - Lista a., b., c. con nombres cortos de tareas
    - Header "Timeline" en verde teal
    - SIN bordes (los bordes solo van en la tabla Gantt)
    - SIN párrafo vacío después (para que Gantt quede pegado)
    """
    num_tasks = len(tasks)

    # Crear tabla: 1 header + n tareas (para Activities) + 1 timeline
    table = doc.add_table(rows=num_tasks + 2, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # NO aplicar bordes a esta tabla - solo la tabla Gantt lleva bordes

    # Fila 0: Activities header
    cell = table.cell(0, 0)
    cell.text = "Activities"
    _set_cell_shading(cell, COLORS['activities_header'])
    for run in cell.paragraphs[0].runs:
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.bold = True

    # Filas 1 a n: Lista de actividades como a., b., c.
    for t_idx, task in enumerate(tasks):
        task_letter = chr(ord('a') + t_idx)
        task_title = task.get('title', '')
        cell = table.cell(t_idx + 1, 0)
        cell.text = f"{task_letter}.  {task_title}"
        _set_cell_shading(cell, COLORS['activities_body'])

    # Última fila: Timeline
    start = phase.get('start_week', phase.get('start_month', ''))
    end = phase.get('end_week', phase.get('end_month', ''))

    cell = table.cell(num_tasks + 1, 0)
    cell.text = "Timeline"
    _set_cell_shading(cell, COLORS['timeline'])
    for run in cell.paragraphs[0].runs:
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.bold = True

    # Aplicar formato Aptos 10 e interlineado sencillo a todas las celdas
    _format_all_table_cells(table)

    # NO añadir párrafo aquí - la tabla Gantt debe quedar pegada


def _add_gantt_table_improved(doc: Document, tasks: list, phase: dict):
    """
    Agrega la tabla Gantt mejorada con:
    - "Weeks" en itálica como header
    - SIEMPRE 12 columnas: 1 para código de tarea + 11 para semanas
    - El rango de semanas se ajusta para cubrir las tareas de la fase
    - Códigos a., b., c. para las tareas
    - Celdas coloreadas (#FFC000) cuando hay actividad
    - E1, E2... para eventos dentro de la celda amarilla
    - D1, D2... para deliverables dentro de la celda amarilla
    - Bordes negros en todas las celdas
    - Leyenda al final alineada a la derecha
    """
    NUM_WEEK_COLUMNS = 11  # Siempre 11 columnas de semanas

    # Determinar el rango de semanas que necesitamos mostrar
    # Buscamos la semana mínima y máxima de todas las tareas
    all_starts = [task.get('start_week', task.get('start_month', 1)) for task in tasks]
    all_ends = [task.get('end_week', task.get('end_month', 1)) for task in tasks]

    min_week = min(all_starts) if all_starts else 1
    max_week = max(all_ends) if all_ends else min_week + 10

    # Calcular el rango de semanas a mostrar (centrado en las tareas si es posible)
    task_span = max_week - min_week + 1

    if task_span <= NUM_WEEK_COLUMNS:
        # Las tareas caben en 11 columnas, empezamos desde min_week
        display_start = min_week
    else:
        # Las tareas no caben, ajustamos para mostrar el rango relevante
        display_start = min_week

    num_tasks = len(tasks)

    # Crear tabla: SIEMPRE 12 columnas (1 tarea + 11 semanas)
    table = doc.add_table(rows=num_tasks + 1, cols=NUM_WEEK_COLUMNS + 1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Aplicar bordes negros a la tabla
    _set_table_borders(table)

    # Fila header con "Weeks" en itálica
    header_row = table.rows[0]
    weeks_cell = header_row.cells[0]
    weeks_cell.text = ""
    p = weeks_cell.paragraphs[0]
    run = p.add_run("Weeks")
    run.italic = True

    # Números de semana (11 columnas)
    for i in range(NUM_WEEK_COLUMNS):
        week_num = display_start + i
        header_row.cells[i + 1].text = str(week_num)

    # Contador de eventos y deliverables para esta fase
    event_counter = 0
    deliverable_counter = 0

    # Filas de tareas
    for t_idx, task in enumerate(tasks):
        row = table.rows[t_idx + 1]

        # Código de tarea (a., b., c., etc.)
        task_letter = chr(ord('a') + t_idx)
        row.cells[0].text = f"{task_letter}."

        # Determinar período de inicio y fin de la tarea
        task_start = task.get('start_week', task.get('start_month', display_start))
        task_end = task.get('end_week', task.get('end_month', task_start))

        # Colorear celdas activas (relativas al display_start)
        for week in range(task_start, task_end + 1):
            col_idx = week - display_start + 1
            if 1 <= col_idx <= NUM_WEEK_COLUMNS:
                cell = row.cells[col_idx]
                _set_cell_shading(cell, COLORS['gantt_active'])

                # Marcar eventos en la celda
                event_week = task.get('event_week', task.get('event_month'))
                if event_week == week:
                    event_counter += 1
                    cell.text = f"E{event_counter}"

                # Marcar deliverables en la celda
                deliverable_week = task.get('deliverable_week', task.get('deliverable_month'))
                if deliverable_week == week:
                    deliverable_counter += 1
                    cell.text = f"D{deliverable_counter}"

    # Aplicar formato Aptos 10 e interlineado sencillo a todas las celdas
    _format_all_table_cells(table)

    # Agregar leyenda como párrafo alineado a la derecha
    legend_para = doc.add_paragraph()
    legend_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = legend_para.add_run("E: Event, D: Deliverable")
    run.italic = True
    run.font.size = Pt(9)


def _add_deliverable_box(doc: Document, deliverable: dict):
    """
    Agrega una caja de entregable con formato del sample:
    - 2 columnas: código (1.73") + nombre (4.17")
    - Colores amarillos
    - Formato Aptos 10
    """
    table = doc.add_table(rows=1, cols=2)

    # Establecer anchos de columnas (en twips: 1 inch = 1440 twips)
    table.columns[0].width = Twips(2494)  # ~1.73 inches
    table.columns[1].width = Twips(6000)  # ~4.17 inches

    # Primera celda: código del deliverable
    cell = table.cell(0, 0)
    cell.width = Twips(2494)
    cell.text = deliverable.get('code', 'Deliverable')
    _set_cell_shading(cell, COLORS['deliverable_header'])
    for run in cell.paragraphs[0].runs:
        run.font.bold = True

    # Segunda celda: nombre del deliverable
    cell = table.cell(0, 1)
    cell.width = Twips(6000)
    cell.text = deliverable.get('name', '')
    _set_cell_shading(cell, COLORS['deliverable_body'])

    # Aplicar formato Aptos 10 a las celdas
    _format_all_table_cells(table)

    doc.add_paragraph()


def _add_experience_box(doc: Document, text: str):
    """Agrega una caja de experiencia previa"""
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = text
    _set_cell_shading(cell, COLORS['experience_box'])
    doc.add_paragraph()
