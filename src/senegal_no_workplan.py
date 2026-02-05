# -*- coding: utf-8 -*-
"""
Metodología P2-P2RS Senegal - Sin WorkPlan para debug
"""

import os
import sys
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from translations import get_language_config


# Importar funciones del document_writer pero sin el WorkPlan
from document_writer import (
    _setup_aninver_styles, _add_paragraph, _add_content_block,
    _add_principles_table, _add_phase_improved
)


def create_word_document_no_workplan(methodology: dict, output_path: str, lang_config: dict):
    """
    Crea un documento Word SIN el WorkPlan consolidado (para debug)
    """
    doc = Document()
    sections = lang_config['sections']

    # Configurar estilos Aninver
    _setup_aninver_styles(doc)

    # Configurar margenes
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # Titulo principal de seccion
    title = doc.add_paragraph(sections['title'], style='Aninver Section Heading')

    # Introduccion si existe
    if methodology.get('introduction'):
        _add_paragraph(doc, methodology['introduction'])

    # Comprension del contexto
    if methodology.get('context'):
        doc.add_paragraph(sections['context'], style='Aninver Title 2')
        _add_content_block(doc, methodology['context'])

    # Principios rectores
    if methodology.get('principles'):
        doc.add_paragraph(sections.get('principles', 'Our guiding principles'), style='Aninver Title 2')
        intro_text = lang_config.get('principles_intro',
            'The following key principles underpin our proposed strategy.')
        _add_paragraph(doc, intro_text)
        _add_principles_table(doc, methodology['principles'], lang_config)

    # Enfoque tecnico y metodologia
    doc.add_paragraph(sections['approach'], style='Aninver Title 1')

    # Fases del proyecto
    if methodology.get('phases'):
        for phase in methodology['phases']:
            _add_phase_improved(doc, phase, lang_config)

    # Gestion de riesgos
    if methodology.get('risks'):
        doc.add_paragraph(sections['risks'], style='Aninver Title 1')
        _add_content_block(doc, methodology['risks'])

    # Aseguramiento de calidad
    if methodology.get('quality'):
        doc.add_paragraph(sections['quality'], style='Aninver Title 1')
        _add_content_block(doc, methodology['quality'])

    # SIN Work Plan consolidado

    # Guardar documento
    doc.save(output_path)


def generate_methodology():
    methodology = {
        "introduction": "Notre consortium propose une approche methodologique integree et participative pour l'accompagnement des communautes des 15 communes cibles du P2-P2RS au Senegal. Forts de notre experience dans les projets de gestion des ressources naturelles au Sahel, nous avons concu une methodologie qui s'ancre dans les realites locales tout en repondant aux exigences de la Banque Africaine de Developpement. Notre approche repose sur trois piliers fondamentaux : la participation communautaire, le renforcement des capacites locales et la durabilite des interventions. Nous mobiliserons une equipe pluridisciplinaire experimentee dans les regions de Matam, Tambacounda et Fatick pour assurer un accompagnement de proximite efficace des Comites de Gestion Concertee des Ressources Naturelles (CGC/RN) et la mise en oeuvre reussie du Mecanisme de Gestion des Plaintes (MGP).",

        "context": """**Le Programme P2RS et son Contexte Regional**

Le Programme regional de resilience a l'insecurite alimentaire et nutritionnelle au Sahel (P2RS), initie par la Banque Africaine de Developpement depuis 2014, constitue une reponse structurelle aux defis chroniques de securite alimentaire dans la region sahelienne. Le P2-P2RS (Projet 2) s'inscrit dans la continuite des acquis de la phase 1, avec un accent particulier sur la consolidation et l'elargissement des interventions dans 15 communes reparties dans les regions de Matam, Tambacounda et Fatick.

**Zones d'Intervention**

Le projet travaille autour de deux poles de developpement. Le Pole Nord (Vallee du fleuve Senegal) regroupe 12 communes : Matam (Agnam Civol, Dabia, Bokidiawe, Nguidjilone), Kanel (Aoure, Orkadjiere, Sinthiou Bamambe, Ndendory) et Bakel (Belle, Sinthiou Fissa, Gabou, Gathiary). Le Pole Centre (Bassin arachidier) comprend 3 communes insulaires (Fimela, Bassoul, Djirnda) dans la region de Fatick.

**Cadre Institutionnel**

Le projet s'inscrit dans le cadre de la politique nationale de decentralisation du Senegal, avec les collectivites territoriales comme porte d'entree principale. Les acteurs cles incluent : les services techniques deconcentres (DRDR, IREF, SAED), les organisations de producteurs, les groupements de femmes et les autorites coutumieres.

**Problematique**

Les communautes cibles font face a plusieurs defis : degradation des terres agricoles, conflits agriculteurs-eleveurs, faible capacite des structures communautaires, et insuffisance des connaissances en GDT et agroecologie.

**Objectifs de la Mission**

L'objectif general est l'appui a l'operationnalisation des CGC/RN, la formation des acteurs a la gestion durable des ressources naturelles, et l'appui a la mise en oeuvre du MGP.""",

        "principles": [
            {
                "name": "Approche Participative",
                "description": "Toutes nos interventions seront fondees sur la participation effective des communautes a toutes les etapes. Nous accorderons une attention particuliere a l'inclusion des groupes vulnerables, notamment les femmes et les jeunes, dans les instances de decision."
            },
            {
                "name": "Ancrage Local",
                "description": "Notre methodologie valorise les connaissances endogenes des communautes en matiere de gestion des ressources naturelles. Les pratiques traditionnelles validees seront integrees aux innovations techniques."
            },
            {
                "name": "Adaptation au Contexte",
                "description": "Nous adapterons nos methodes aux specificites de chaque zone agro-ecologique. Les outils de communication seront developpes dans les langues locales (Pulaar, Wolof, Serere)."
            },
            {
                "name": "Apprentissage par l'Action",
                "description": "Notre programme privilegie l'apprentissage pratique a travers les Champs Ecoles Paysans et les demonstrations sur le terrain."
            },
            {
                "name": "Coordination",
                "description": "Nous assurerons une coordination etroite avec les services techniques, les partenaires du P2RS et les projets connexes."
            },
            {
                "name": "Equite Genre",
                "description": "L'integration du genre sera transversale. Les femmes representeront au moins 40% des participants aux formations et 30% des membres des CGC/RN."
            }
        ],

        "phases": [
            {
                "title": "Phase 1 : Demarrage et Diagnostic Participatif",
                "description": "Cette phase de 4 mois etablira les bases de notre intervention a travers une comprehension approfondie du contexte socioeconomique des 15 communes. Nous installerons nos equipes dans les trois regions et conduirons le diagnostic participatif selon une approche multi-acteurs.",
                "start_week": 1,
                "end_week": 16,
                "tasks": [
                    {
                        "code": "1A",
                        "title": "Installation et cadrage methodologique",
                        "description": "Installation des bases operationnelles, recrutement des 13 animateurs, mise en place des moyens logistiques et atelier de lancement.",
                        "items": [
                            "Installation des bases a Matam, Bakel et Fatick",
                            "Recrutement et formation des animateurs",
                            "Atelier de lancement avec les parties prenantes",
                            "Visites aux autorites des 15 communes"
                        ],
                        "start_week": 1,
                        "end_week": 4,
                        "deliverable_week": 4,
                        "deliverable_code": "L1"
                    },
                    {
                        "code": "1B",
                        "title": "Etudes diagnostiques socioeconomiques",
                        "description": "Diagnostic participatif combinant enquetes menages, focus groups et cartographie des ressources naturelles.",
                        "items": [
                            "Elaboration des outils de collecte",
                            "Enquetes et focus groups dans les 15 communes",
                            "Cartographie participative",
                            "Analyse des donnees"
                        ],
                        "start_week": 3,
                        "end_week": 12
                    },
                    {
                        "code": "1C",
                        "title": "Restitution et validation",
                        "description": "Restitution des resultats lors d'ateliers communaux et regionaux.",
                        "items": [
                            "15 ateliers communaux",
                            "3 ateliers regionaux",
                            "Finalisation des rapports"
                        ],
                        "start_week": 13,
                        "end_week": 16,
                        "deliverable_week": 16,
                        "deliverable_code": "L2"
                    }
                ],
                "deliverables": [
                    {"code": "L1", "name": "Rapport d'etablissement"},
                    {"code": "L2", "name": "Rapports de diagnostic des 15 communes"}
                ]
            },
            {
                "title": "Phase 2 : Mise en Place des CGC/RN",
                "description": "Cette phase de 6 mois etablira un systeme durable de gestion concertee des ressources naturelles a travers la mise en place des CGC/RN et la negociation des conventions locales.",
                "start_week": 17,
                "end_week": 40,
                "tasks": [
                    {
                        "code": "2A",
                        "title": "Sensibilisation communautaire",
                        "description": "Campagne d'information dans les 15 communes utilisant reunions, radio et supports visuels.",
                        "items": [
                            "Strategie de communication",
                            "Emissions radio en langues locales",
                            "Reunions villageoises",
                            "Panneaux d'information"
                        ],
                        "start_week": 17,
                        "end_week": 24
                    },
                    {
                        "code": "2B",
                        "title": "Constitution des CGC/RN",
                        "description": "Mise en place des comites avec representation inclusive de tous les groupes d'utilisateurs.",
                        "items": [
                            "Evaluation des structures existantes",
                            "Assemblees constitutives",
                            "Election des bureaux",
                            "Elaboration des statuts"
                        ],
                        "start_week": 22,
                        "end_week": 32
                    },
                    {
                        "code": "2C",
                        "title": "Conventions locales et delimitation",
                        "description": "Negociation des regles de gestion et materialisation des espaces.",
                        "items": [
                            "Ateliers de negociation",
                            "Redaction des conventions",
                            "Adoption par conseils municipaux",
                            "Bornage des couloirs de passage"
                        ],
                        "start_week": 28,
                        "end_week": 40,
                        "deliverable_week": 40,
                        "deliverable_code": "L3"
                    }
                ],
                "deliverables": [
                    {"code": "L3", "name": "15 CGC/RN operationnels avec conventions locales"}
                ]
            },
            {
                "title": "Phase 3 : Formation et Renforcement des Capacites",
                "description": "Phase de 12 mois dediee au renforcement des capacites en GDT, agroecologie, agroforesterie et RNA a travers les Champs Ecoles Paysans.",
                "start_week": 24,
                "end_week": 72,
                "tasks": [
                    {
                        "code": "3A",
                        "title": "Plan de renforcement des capacites",
                        "description": "Elaboration du plan detaille definissant thematiques, approches et calendrier.",
                        "items": [
                            "Analyse des besoins",
                            "Definition des modules",
                            "Calendrier de formation"
                        ],
                        "start_week": 24,
                        "end_week": 28,
                        "deliverable_week": 28,
                        "deliverable_code": "L4"
                    },
                    {
                        "code": "3B",
                        "title": "Formation des formateurs",
                        "description": "Sessions pour techniciens et animateurs sur facilitation, GDT, agroecologie.",
                        "items": [
                            "Techniques de facilitation (5 jours)",
                            "GDT et CES/DRS (5 jours)",
                            "Agroecologie et RNA (5 jours)",
                            "Gestion des conflits (3 jours)"
                        ],
                        "start_week": 28,
                        "end_week": 36
                    },
                    {
                        "code": "3C",
                        "title": "Formation des CGC/RN et producteurs",
                        "description": "Formation des membres des CGC/RN et deploiement de 30 Champs Ecoles Paysans.",
                        "items": [
                            "Modules CGC/RN (gouvernance, GDT)",
                            "Installation 30 CEP",
                            "Sessions pratiques",
                            "Evaluation des acquis"
                        ],
                        "start_week": 36,
                        "end_week": 72,
                        "deliverable_week": 72,
                        "deliverable_code": "L5"
                    }
                ],
                "deliverables": [
                    {"code": "L4", "name": "Plan de renforcement des capacites"},
                    {"code": "L5", "name": "Rapport de formation (450 membres CGC/RN, 900 producteurs)"}
                ]
            },
            {
                "title": "Phase 4 : Accompagnement et MGP",
                "description": "Phase transversale de 20 mois pour l'accompagnement des CGC/RN et l'operationnalisation du Mecanisme de Gestion des Plaintes.",
                "start_week": 20,
                "end_week": 96,
                "tasks": [
                    {
                        "code": "4A",
                        "title": "Installation du MGP",
                        "description": "Mise en place des comites de gestion des plaintes aux niveaux villageois, communal et regional.",
                        "items": [
                            "Identification des membres",
                            "Installation des comites",
                            "Points de depot et numero vert",
                            "Procedures de traitement"
                        ],
                        "start_week": 20,
                        "end_week": 32
                    },
                    {
                        "code": "4B",
                        "title": "Formation et vulgarisation MGP",
                        "description": "Formation des membres des comites et campagne de vulgarisation.",
                        "items": [
                            "Formation sur procedures MGP",
                            "Techniques de mediation",
                            "Spots radio",
                            "Reunions d'information"
                        ],
                        "start_week": 28,
                        "end_week": 40,
                        "deliverable_week": 40,
                        "deliverable_code": "L6"
                    },
                    {
                        "code": "4C",
                        "title": "Accompagnement et suivi-evaluation",
                        "description": "Accompagnement continu des CGC/RN et capitalisation.",
                        "items": [
                            "Visites hebdomadaires",
                            "Suivi des conventions",
                            "Rapports trimestriels",
                            "Capitalisation finale"
                        ],
                        "start_week": 40,
                        "end_week": 96,
                        "deliverable_week": 96,
                        "deliverable_code": "L7"
                    }
                ],
                "deliverables": [
                    {"code": "L6", "name": "Dispositif MGP operationnel"},
                    {"code": "L7", "name": "8 rapports trimestriels et rapport final"}
                ]
            }
        ],

        "risks": """**Risques Contextuels**

- Instabilite securitaire : Coordination avec autorites locales et adaptation des plannings.
- Conditions climatiques : Flexibilite du calendrier.
- Conflits intercommunautaires : Mediation preventive.

**Risques Operationnels**

- Mobilisation insuffisante : Approche participative et communication intensive.
- Turnover du personnel : Documentation et partage des connaissances.
- Difficultes logistiques : Moyens adaptes (pirogues pour les iles).

**Risques Institutionnels**

- Faible appropriation locale : Implication des maires des le depart.
- Conflits de competences : Coordination reguliere avec partenaires.""",

        "quality": """**Indicateurs Cles de Performance**

- Taux de fonctionnement des CGC/RN : cible 80%
- Taux de resolution des plaintes : cible 90%
- Satisfaction des beneficiaires : cible 85%
- Representation des femmes : minimum 30%

**Mecanismes de Controle**

- Supervision mensuelle du chef de mission
- Reunions hebdomadaires de coordination
- Visites conjointes avec l'UCP (trimestrielles)
- Evaluations participatives semestrielles

**Protocoles de Reporting**

- Rapports hebdomadaires des animateurs
- Rapports mensuels des techniciens
- Rapports trimestriels a l'UCP"""
    }

    return methodology


if __name__ == "__main__":
    methodology = generate_methodology()
    lang_config = get_language_config('fr')
    output_path = "/Users/apple/metodologias-rfps/output/Methodologie_P2RS_Senegal_noworkplan.docx"
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    create_word_document_no_workplan(methodology, output_path, lang_config)
    print(f"Document genere: {output_path}")
