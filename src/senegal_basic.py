#!/usr/bin/env python3
"""
Generador básico de metodología para P2-P2RS Senegal
Sin estilos personalizados - solo estilos nativos de Word
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_document():
    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # ===== TÍTULO =====
    p = doc.add_paragraph()
    run = p.add_run("SECTION D : DESCRIPTION DE L'APPROCHE, MÉTHODOLOGIE ET PLAN DE TRAVAIL")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ===== INTRODUCTION =====
    p = doc.add_paragraph()
    run = p.add_run("INTRODUCTION")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    doc.add_paragraph("""Notre consortium propose une approche méthodologique intégrée et participative pour l'accompagnement des communautés des 15 communes cibles du P2-P2RS au Sénégal. Forts de notre expérience dans les projets de gestion des ressources naturelles au Sahel, nous avons conçu une méthodologie qui s'ancre dans les réalités locales tout en répondant aux exigences de la Banque Africaine de Développement.

Notre approche repose sur trois piliers fondamentaux : la participation communautaire, le renforcement des capacités locales et la durabilité des interventions. Nous mobiliserons une équipe pluridisciplinaire expérimentée dans les régions de Matam, Tambacounda et Fatick pour assurer un accompagnement de proximité efficace des Comités de Gestion Concertée des Ressources Naturelles (CGC/RN) et la mise en œuvre réussie du Mécanisme de Gestion des Plaintes (MGP).""")

    # ===== COMPRÉHENSION DU CONTEXTE =====
    p = doc.add_paragraph()
    run = p.add_run("1. COMPRÉHENSION DU CONTEXTE")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    p = doc.add_paragraph()
    run = p.add_run("1.1 Contexte du Programme P2RS")
    run.bold = True

    doc.add_paragraph("""Le Programme régional de résilience à l'insécurité alimentaire et nutritionnelle au Sahel (P2RS), initié par la Banque Africaine de Développement depuis 2014, constitue une réponse structurelle aux défis chroniques de sécurité alimentaire dans la région sahélienne. Le P2-P2RS (Projet 2) s'inscrit dans la continuité des acquis de la phase 1, avec un accent particulier sur la consolidation et l'élargissement des interventions dans 15 communes réparties dans les régions de Matam, Tambacounda et Fatick.""")

    p = doc.add_paragraph()
    run = p.add_run("1.2 Zones d'intervention")
    run.bold = True

    doc.add_paragraph("""Le projet intervient dans deux pôles de développement distincts :

Le Pôle Nord (Vallée du fleuve Sénégal) regroupe 12 communes :
- Région de Tambacounda, Département de Bakel : Bellé, Sinthiou Fissa, Gabou, Gathiary
- Région de Matam, Département de Matam : Agnam Civol, Dabia, Bokidiawé, Nguidjilone
- Région de Matam, Département de Kanel : Aouré, Orkadjiéré, Sinthiou Bamambé, Ndendory

Le Pôle Centre (Bassin arachidier) comprend 3 communes insulaires :
- Région de Fatick : Fimela, Bassoul, Djirnda

Ces communes sont caractérisées par une vulnérabilité aux changements climatiques, des conflits d'usage entre agriculture et élevage, et une pression sur les ressources naturelles.""")

    p = doc.add_paragraph()
    run = p.add_run("1.3 Objectifs de la mission")
    run.bold = True

    doc.add_paragraph("""L'objectif général est l'appui à l'opérationnalisation des Comités de Gestion Concertée des Ressources Naturelles (CGC/RN), la formation des acteurs à la gestion durable des ressources naturelles et des bonnes pratiques de GDT/agroécologie, et l'appui à la mise en œuvre du Mécanisme de Gestion des Plaintes (MGP).""")

    # ===== PRINCIPES DIRECTEURS =====
    p = doc.add_paragraph()
    run = p.add_run("2. NOS PRINCIPES DIRECTEURS")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    doc.add_paragraph("Les principes clés suivants sous-tendent notre stratégie proposée :")

    principles = [
        ("Principe 1 : Approche participative et inclusive",
         "Toutes nos interventions seront fondées sur la participation effective des communautés à toutes les étapes : diagnostic, planification, mise en œuvre et suivi. Nous accorderons une attention particulière à l'inclusion des groupes vulnérables, notamment les femmes et les jeunes."),

        ("Principe 2 : Ancrage dans les savoirs locaux",
         "Notre méthodologie valorise les connaissances endogènes des communautés en matière de gestion des ressources naturelles. Les pratiques traditionnelles validées seront intégrées aux innovations techniques."),

        ("Principe 3 : Adaptation au contexte local",
         "Nous adapterons nos méthodes aux spécificités de chaque zone agro-écologique. Les outils de communication seront développés dans les langues locales (Pulaar, Wolof, Sérère)."),

        ("Principe 4 : Apprentissage par l'action",
         "Notre programme privilégie l'apprentissage pratique à travers les Champs Écoles Paysans et les démonstrations sur le terrain."),

        ("Principe 5 : Coordination et synergie",
         "Nous assurerons une coordination étroite avec les services techniques, les partenaires du P2RS et les projets connexes."),

        ("Principe 6 : Équité genre",
         "L'intégration du genre sera transversale. Les femmes représenteront au moins 40% des participants aux formations et 30% des membres des CGC/RN.")
    ]

    for name, desc in principles:
        p = doc.add_paragraph()
        run = p.add_run(name)
        run.bold = True
        doc.add_paragraph(desc)

    # ===== APPROCHE MÉTHODOLOGIQUE =====
    p = doc.add_paragraph()
    run = p.add_run("3. APPROCHE TECHNIQUE ET MÉTHODOLOGIE")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    doc.add_paragraph("Notre intervention est structurée en 4 phases complémentaires sur 24 mois :")

    # PHASE 1
    p = doc.add_paragraph()
    run = p.add_run("3.1 Phase 1 : Démarrage et Diagnostic Participatif (Mois 1-4)")
    run.bold = True
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    doc.add_paragraph("""Cette phase fondamentale de 4 mois permettra d'établir les bases solides de notre intervention à travers une compréhension approfondie du contexte socioéconomique et des dynamiques de gestion des ressources naturelles dans les 15 communes cibles.""")

    p = doc.add_paragraph()
    run = p.add_run("Activité 1A : Installation et cadrage méthodologique")
    run.bold = True
    run.italic = True

    doc.add_paragraph("""L'activité de démarrage comprendra l'installation des bases opérationnelles dans les trois régions d'intervention :
- Installation des bases opérationnelles à Matam, Bakel et Fatick
- Recrutement et formation des 13 animateurs locaux
- Acquisition et déploiement des moyens logistiques
- Atelier de lancement avec les parties prenantes du P2RS
- Visites de courtoisie aux autorités des 15 communes""")

    p = doc.add_paragraph()
    run = p.add_run("Livrable L1 : Rapport d'établissement avec recadrage méthodologique et planning actualisé")
    run.bold = True

    p = doc.add_paragraph()
    run = p.add_run("Activité 1B : Études diagnostiques socioéconomiques")
    run.bold = True
    run.italic = True

    doc.add_paragraph("""Le diagnostic socioéconomique sera conduit dans chacun des villages ciblés :
- Élaboration des outils de collecte de données
- Formation des enquêteurs sur les méthodologies participatives
- Réalisation des enquêtes ménages et focus groups dans les 15 communes
- Cartographie participative des ressources naturelles
- Analyse des données et rédaction des monographies communales""")

    p = doc.add_paragraph()
    run = p.add_run("Activité 1C : Restitution et validation du diagnostic")
    run.bold = True
    run.italic = True

    doc.add_paragraph("""Les résultats seront restitués et validés :
- Organisation de 15 ateliers communaux de restitution
- Organisation de 3 ateliers régionaux de validation
- Finalisation des rapports de diagnostic par région""")

    p = doc.add_paragraph()
    run = p.add_run("Livrable L2 : Rapports de diagnostic socioéconomique et environnemental des 15 communes")
    run.bold = True

    # PHASE 2
    p = doc.add_paragraph()
    run = p.add_run("3.2 Phase 2 : Mise en Place et Opérationnalisation des CGC/RN (Mois 5-10)")
    run.bold = True
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    doc.add_paragraph("""Cette phase de 6 mois constitue le cœur de notre intervention pour l'établissement d'un système durable de gestion concertée des ressources naturelles.""")

    p = doc.add_paragraph()
    run = p.add_run("Activité 2A : Sensibilisation et mobilisation communautaire")
    run.bold = True
    run.italic = True

    doc.add_paragraph("""Campagne d'information et de sensibilisation :
- Élaboration de la stratégie et des supports de communication
- Production d'émissions radiophoniques en Pulaar, Wolof et Sérère
- Organisation de réunions d'information dans chaque village
- Formation de relais communautaires
- Mise en place de panneaux d'information""")

    p = doc.add_paragraph()
    run = p.add_run("Activité 2B : Constitution et structuration des CGC/RN")
    run.bold = True
    run.italic = True

    doc.add_paragraph("""Mise en place des comités :
- Évaluation des structures existantes de gestion des RN
- Facilitation des assemblées générales constitutives
- Élection des bureaux des CGC/RN (avec quota genre)
- Élaboration des statuts et règlements intérieurs
- Reconnaissance officielle des CGC/RN par les autorités""")

    p = doc.add_paragraph()
    run = p.add_run("Activité 2C : Négociation et élaboration des conventions locales")
    run.bold = True
    run.italic = True

    doc.add_paragraph("""Négociation des règles de gestion :
- Organisation d'ateliers de négociation multi-acteurs par commune
- Facilitation des consensus sur les règles de gestion
- Rédaction et validation des conventions locales
- Adoption des conventions par les conseils municipaux
- Délimitation et matérialisation des espaces (couloirs de passage, zones de pâturage)""")

    p = doc.add_paragraph()
    run = p.add_run("Livrable L3 : 15 CGC/RN opérationnels avec conventions locales adoptées")
    run.bold = True

    # PHASE 3
    p = doc.add_paragraph()
    run = p.add_run("3.3 Phase 3 : Formation et Renforcement des Capacités (Mois 6-18)")
    run.bold = True
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    doc.add_paragraph("""Cette phase de 12 mois est dédiée au renforcement des capacités en GDT, agroécologie, agroforesterie et RNA.""")

    p = doc.add_paragraph()
    run = p.add_run("Activité 3A : Élaboration du plan de renforcement des capacités")
    run.bold = True
    run.italic = True

    doc.add_paragraph("""Plan détaillé de formation :
- Analyse des besoins de formation par groupe cible
- Définition des objectifs et contenus pédagogiques
- Identification des formateurs et personnes ressources
- Élaboration du calendrier de formation""")

    p = doc.add_paragraph()
    run = p.add_run("Livrable L4 : Plan de renforcement des capacités validé")
    run.bold = True

    p = doc.add_paragraph()
    run = p.add_run("Activité 3B : Formation des formateurs et animateurs")
    run.bold = True
    run.italic = True

    doc.add_paragraph("""Sessions de formation :
- Session sur les techniques de facilitation (5 jours)
- Session technique sur la GDT et CES/DRS (5 jours)
- Session technique sur l'agroécologie et la RNA (5 jours)
- Session sur la production de plants agroforestiers (3 jours)
- Session sur la gestion des conflits et médiation (3 jours)""")

    p = doc.add_paragraph()
    run = p.add_run("Activité 3C : Formation des CGC/RN et producteurs")
    run.bold = True
    run.italic = True

    doc.add_paragraph("""Formation des bénéficiaires :
- Module gouvernance et leadership (2 jours x 15 CGC/RN)
- Module techniques de GDT (3 jours x 15 CGC/RN)
- Module gestion des conflits (2 jours x 15 CGC/RN)
- Installation de 30 Champs Écoles Paysans (2 par commune)
- Sessions pratiques sur le compostage, ouvrages antiérosifs, RNA""")

    p = doc.add_paragraph()
    run = p.add_run("Livrable L5 : Rapport de formation avec 450 membres CGC/RN et 900 producteurs formés")
    run.bold = True

    # PHASE 4
    p = doc.add_paragraph()
    run = p.add_run("3.4 Phase 4 : Accompagnement et Mise en Œuvre du MGP (Mois 5-24)")
    run.bold = True
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    doc.add_paragraph("""Cette phase transversale de 20 mois est dédiée à l'accompagnement des CGC/RN et à l'opérationnalisation du Mécanisme de Gestion des Plaintes.""")

    p = doc.add_paragraph()
    run = p.add_run("Activité 4A : Installation du dispositif MGP")
    run.bold = True
    run.italic = True

    doc.add_paragraph("""Mise en place du mécanisme :
- Identification des membres des comités de gestion des plaintes
- Organisation des assemblées d'installation des comités
- Dotation en matériels et outils de travail
- Mise en place des points de dépôt et du numéro vert
- Élaboration des procédures et circuits de traitement""")

    p = doc.add_paragraph()
    run = p.add_run("Activité 4B : Formation et vulgarisation du MGP")
    run.bold = True
    run.italic = True

    doc.add_paragraph("""Formation et sensibilisation :
- Formation sur la politique et les procédures du MGP (2 jours)
- Formation sur les techniques de médiation (2 jours)
- Production de spots radio sur le MGP en langues locales
- Organisation de réunions d'information dans tous les villages""")

    p = doc.add_paragraph()
    run = p.add_run("Livrable L6 : Dispositif MGP opérationnel avec comités formés")
    run.bold = True

    p = doc.add_paragraph()
    run = p.add_run("Activité 4C : Accompagnement continu et suivi-évaluation")
    run.bold = True
    run.italic = True

    doc.add_paragraph("""Accompagnement des CGC/RN :
- Visites hebdomadaires d'accompagnement des CGC/RN
- Appui à l'animation des réunions statutaires
- Suivi de l'application des conventions locales
- Facilitation de la résolution des conflits
- Élaboration des rapports trimestriels d'activités
- Capitalisation et documentation des bonnes pratiques""")

    p = doc.add_paragraph()
    run = p.add_run("Livrable L7 : 8 rapports trimestriels d'activités et rapport final de capitalisation")
    run.bold = True

    # ===== GESTION DES RISQUES =====
    p = doc.add_paragraph()
    run = p.add_run("4. GESTION DES RISQUES")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    p = doc.add_paragraph()
    run = p.add_run("4.1 Risques contextuels")
    run.bold = True

    doc.add_paragraph("""- Instabilité sécuritaire : La région nord peut connaître des tensions. Mesure d'atténuation : coordination étroite avec les autorités locales et adaptation des plannings.
- Conditions climatiques défavorables : Les aléas peuvent perturber les activités. Mesure : flexibilité du calendrier.
- Conflits intercommunautaires : Tensions agriculteurs-éleveurs possibles. Mesure : médiation préventive et implication des autorités coutumières.""")

    p = doc.add_paragraph()
    run = p.add_run("4.2 Risques opérationnels")
    run.bold = True

    doc.add_paragraph("""- Mobilisation insuffisante des communautés : Mesure : approche participative et communication intensive.
- Turnover du personnel : Mesure : documentation systématique et partage des connaissances.
- Difficultés logistiques : Mesure : planification anticipée et moyens adaptés (pirogues pour les îles).""")

    p = doc.add_paragraph()
    run = p.add_run("4.3 Risques institutionnels")
    run.bold = True

    doc.add_paragraph("""- Faible appropriation par les autorités locales : Mesure : implication dès le départ des maires et conseils municipaux.
- Conflits de compétences : Mesure : cartographie des acteurs et coordination régulière.""")

    # ===== ASSURANCE QUALITÉ =====
    p = doc.add_paragraph()
    run = p.add_run("5. ASSURANCE QUALITÉ")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    p = doc.add_paragraph()
    run = p.add_run("5.1 Indicateurs clés de performance (KPIs)")
    run.bold = True

    doc.add_paragraph("""- Taux de fonctionnement des CGC/RN (réunions régulières) : cible 80%
- Taux de résolution des plaintes dans les délais : cible 90%
- Taux de satisfaction des bénéficiaires des formations : cible 85%
- Représentation des femmes dans les instances de décision : minimum 30%
- Taux d'application des conventions locales : cible 75%""")

    p = doc.add_paragraph()
    run = p.add_run("5.2 Mécanismes de contrôle et supervision")
    run.bold = True

    doc.add_paragraph("""- Supervision mensuelle du chef de mission dans chaque région
- Réunions hebdomadaires de coordination des techniciens
- Visites de terrain conjointes avec l'UCP (trimestrielles)
- Audits externes annuels des activités
- Évaluations participatives semestrielles avec les communautés""")

    p = doc.add_paragraph()
    run = p.add_run("5.3 Protocoles de reporting")
    run.bold = True

    doc.add_paragraph("""- Rapports hebdomadaires des animateurs (format standardisé)
- Rapports mensuels des techniciens avec données de suivi
- Rapports trimestriels consolidés à l'UCP (5 jours après fin trimestre)
- Base de données géoréférencée des interventions""")

    # ===== PLAN DE TRAVAIL =====
    p = doc.add_paragraph()
    run = p.add_run("6. PLAN DE TRAVAIL")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    doc.add_paragraph("Le tableau ci-dessous présente le chronogramme synthétique des activités sur 24 mois :")

    # Tabla simple
    table = doc.add_table(rows=10, cols=6)
    table.style = 'Table Grid'

    # Headers
    headers = ['Phase / Activité', 'M1-4', 'M5-8', 'M9-12', 'M13-18', 'M19-24']
    for i, h in enumerate(headers):
        table.cell(0, i).text = h

    # Datos
    data = [
        ['Phase 1: Diagnostic', 'XXX', '', '', '', ''],
        ['  L1: Rapport établissement', 'X', '', '', '', ''],
        ['  L2: Rapports diagnostic', 'X', '', '', '', ''],
        ['Phase 2: CGC/RN', '', 'XXX', 'X', '', ''],
        ['  L3: CGC/RN opérationnels', '', '', 'X', '', ''],
        ['Phase 3: Formation', '', 'XX', 'XXX', 'XXX', ''],
        ['  L4-L5: Plan et rapports', '', '', 'X', 'X', ''],
        ['Phase 4: MGP', '', 'XX', 'XXX', 'XXX', 'XXX'],
        ['  L6-L7: MGP et rapports', '', '', 'X', 'X', 'X'],
    ]

    for i, row_data in enumerate(data):
        for j, cell_data in enumerate(row_data):
            table.cell(i+1, j).text = cell_data

    # Guardar
    output_path = "/Users/apple/metodologias-rfps/output/Methodologie_P2RS_Senegal.docx"
    doc.save(output_path)
    print(f"Document généré: {output_path}")


if __name__ == "__main__":
    create_document()
