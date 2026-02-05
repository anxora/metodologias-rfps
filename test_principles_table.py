"""
Test script for Guiding Principles table only
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# Colores Aninver
COLORS = {
    'principle_header': 'BDD6EE',      # Azul claro - header de principios
    'principle_title': '2E74B5',        # Azul oscuro - título de principio
}


def _set_cell_shading(cell, color_hex: str):
    """Establece el color de fondo de una celda"""
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def _remove_table_borders(table):
    """Elimina todos los bordes de una tabla"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')

    tblBorders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )

    existing_borders = tblPr.find(qn('w:tblBorders'))
    if existing_borders is not None:
        tblPr.remove(existing_borders)

    tblPr.append(tblBorders)

    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def create_test_document():
    doc = Document()

    # Configurar márgenes
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # Título
    p = doc.add_paragraph("Our Guiding Principles")
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(11)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    # Intro
    intro = doc.add_paragraph("The following key principles underpin our proposed strategy, and stem from our analysis of the local context as well as the specific needs expressed in the ToR.")
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Datos de prueba - 6 principios
    principles = [
        {
            "name": "AfDB Compliance",
            "description": "Adherence to AfDB procurement guidelines and safeguards is paramount, ensuring all deliverables like RFPs, BOQs, and financial models incorporate eligibility criteria, ESIA/ESMP, anti-corruption measures, and transparency. For this iDICE project, this principle guides NFI upgrade procurement documents and PPP studio tendering."
        },
        {
            "name": "Stakeholder Inclusivity",
            "description": "Extensive consultations with NFI, NFC, Federal Ministry, Nollywood guilds, and regional stakeholders ensure buy-in for NFI action plans and PPP structuring. Applied specifically, this involves 20+ workshops validating scopes for Lagos mega-studio and 100 mini-studios, incorporating youth, gender, and SME voices."
        },
        {
            "name": "Value-for-Money Optimization",
            "description": "VfM analysis using public sector comparator (PSC) models evaluates PPP viability against traditional procurement for film studios, assessing full cycle costs, risk transfer, and NPV/IRR. Tailored to ToR, this determines optimal delivery for infrastructure like soundstages, ensuring affordability limits under AfDB financing and sustainable revenue from studio operations."
        },
        {
            "name": "Technical Excellence",
            "description": "Drawing on global best practices in film infrastructure and creative industry development, the consultants will ensure all technical specifications, designs, and capacity building programs meet international standards while being adapted to Nigeria's specific context and the Nollywood industry's unique requirements."
        },
        {
            "name": "Financial Sustainability",
            "description": "All proposed solutions will be designed with long-term financial viability in mind, incorporating robust revenue models, realistic cost projections, and risk mitigation strategies that ensure the NFI and PPP studios can operate sustainably beyond the project period without continuous donor support."
        },
        {
            "name": "Knowledge Transfer",
            "description": "The team at Aninver will prioritize building local capacity throughout the engagement, ensuring that BOI, NFI, and other stakeholders develop the skills and knowledge needed to manage upgraded facilities, implement PPP agreements, and sustain institutional improvements independently."
        }
    ]

    # Crear tabla de principios
    _add_principles_table(doc, principles)

    # Guardar
    output_path = "/Users/apple/Downloads/Test_Principles_Table.docx"
    doc.save(output_path)
    print(f"Documento guardado en: {output_path}")


def _add_principles_table(doc: Document, principles: list):
    """
    Tabla de principios con 5 columnas:
    - Columnas 0, 2, 4: Contenido (ocupan casi todo el ancho)
    - Columnas 1, 3: Separadores de 1 cm
    """
    if not principles:
        return

    num_principles = len(principles)
    num_rows = 3 if num_principles <= 3 else 5

    # Crear tabla con 5 columnas
    table = doc.add_table(rows=num_rows, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False  # Desactivar autofit para controlar anchos

    # Eliminar bordes
    _remove_table_borders(table)

    # Ancho total disponible: 6 inches (página 8.5" - márgenes 1.25" x 2 = 6")
    # 6 inches = 8640 twips (1 inch = 1440 twips)
    # Separadores: 2 x 1cm = 2 x 567 twips = 1134 twips
    # Contenido disponible: 8640 - 1134 = 7506 twips para 3 columnas = 2502 twips cada una

    separator_width = Cm(1)  # 1 cm para separadores
    content_width = Cm(5.2)  # ~5.2 cm cada columna de contenido (total ~15.6cm + 2cm = 17.6cm ≈ 6")

    # Definir anchos por índice de columna
    col_widths = [content_width, separator_width, content_width, separator_width, content_width]

    # Aplicar anchos a cada celda individualmente (python-docx requiere esto)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = col_widths[idx]

    # Mapeo de columnas de contenido
    content_cols = [0, 2, 4]

    # === FILA 0: Headers "Principle 1", "Principle 2", "Principle 3" ===
    row0 = table.rows[0]
    for i, col_idx in enumerate(content_cols[:min(num_principles, 3)]):
        cell = row0.cells[col_idx]
        cell.text = f"Principle {i + 1}"
        _set_cell_shading(cell, COLORS['principle_header'])
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.name = 'Aptos'
                run.font.size = Pt(10)

    # === FILA 1: Nombres de principios 1-3 ===
    row1 = table.rows[1]
    for i, principle in enumerate(principles[:3]):
        col_idx = content_cols[i]
        cell = row1.cells[col_idx]
        cell.text = principle.get('name', '')
        _set_cell_shading(cell, COLORS['principle_title'])
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.name = 'Aptos'
                run.font.size = Pt(10)

    # === FILA 2: Descripciones de principios 1-3 ===
    row2 = table.rows[2]
    for i, principle in enumerate(principles[:3]):
        col_idx = content_cols[i]
        cell = row2.cells[col_idx]
        cell.text = principle.get('description', '')
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in para.runs:
                run.font.name = 'Aptos'
                run.font.size = Pt(10)

    # === Si hay más de 3 principios ===
    if num_principles > 3:
        # FILA 3: Nombres de principios 4-6
        row3 = table.rows[3]
        for i, principle in enumerate(principles[3:6]):
            col_idx = content_cols[i]
            cell = row3.cells[col_idx]
            cell.text = principle.get('name', '')
            _set_cell_shading(cell, COLORS['principle_title'])
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.bold = True
                    run.font.name = 'Aptos'
                    run.font.size = Pt(10)

        # FILA 4: Descripciones de principios 4-6
        row4 = table.rows[4]
        for i, principle in enumerate(principles[3:6]):
            col_idx = content_cols[i]
            cell = row4.cells[col_idx]
            cell.text = principle.get('description', '')
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in para.runs:
                    run.font.name = 'Aptos'
                    run.font.size = Pt(10)

    # Aplicar interlineado sencillo a todas las celdas
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.line_spacing = 1.0
                para.paragraph_format.space_before = Pt(3)
                para.paragraph_format.space_after = Pt(3)

    doc.add_paragraph()


if __name__ == "__main__":
    create_test_document()
